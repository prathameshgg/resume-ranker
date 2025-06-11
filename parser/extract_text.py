import PyPDF2
from docx import Document
import os
import logging

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Extracting from a PDF
def extract_text_from_pdf(file_path):
    """Extract text from a PDF file."""
    try:
        # Check if file exists
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"PDF file not found: {file_path}")
            
        # Check file size
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            raise ValueError("PDF file is empty")
            
        logger.info(f"Processing PDF file: {file_path} (size: {file_size} bytes)")
        
        with open(file_path, 'rb') as file:
            try:
                reader = PyPDF2.PdfReader(file)
                
                # Check if PDF is encrypted
                if reader.is_encrypted:
                    raise ValueError("PDF is password-protected and cannot be processed")
                
                # Check if PDF has pages
                if len(reader.pages) == 0:
                    raise ValueError("PDF has no pages")
                
                text = ""
                for i, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                        else:
                            logger.warning(f"Page {i+1} extracted no text - might be an image or scanned page")
                    except Exception as page_error:
                        logger.warning(f"Error extracting text from page {i+1}: {str(page_error)}")
                
                if not text.strip():
                    raise ValueError("No text could be extracted from the PDF - it might be a scanned document")
                
                logger.info(f"Successfully extracted {len(text)} characters from PDF")
                return text
            except Exception as pdf_error:
                raise Exception(f"Error reading PDF: {str(pdf_error)}")
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        raise Exception(f"Error extracting text from PDF: {str(e)}")

# Extracting from a Word docx
def extract_text_from_docx(file_path):
    """Extract text from a DOCX file."""
    try:
        # Check if file exists
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
            
        # Check file size
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            raise ValueError("DOCX file is empty")
            
        logger.info(f"Processing DOCX file: {file_path} (size: {file_size} bytes)")
        
        doc = Document(file_path)
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
        
        # Extract text from headers and footers
        for section in doc.sections:
            for header in section.header.paragraphs:
                if header.text.strip():
                    text += header.text + "\n"
            for footer in section.footer.paragraphs:
                if footer.text.strip():
                    text += footer.text + "\n"
        
        # Clean up the text
        text = text.strip()
        
        if not text:
            raise ValueError("No text could be extracted from the DOCX file")
            
        logger.info(f"Successfully extracted {len(text)} characters from DOCX")
        return text
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        raise Exception(f"Error extracting text from DOCX: {str(e)}")